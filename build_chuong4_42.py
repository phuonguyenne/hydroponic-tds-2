# -*- coding: utf-8 -*-
"""Generate figures and insert section 4.2 body + images into CHƯƠNG 4.docx."""
from __future__ import annotations

import os
import sys
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib import font_manager
from docx import Document
from docx.shared import Mm

BASE = Path(r"c:\xampp\htdocs\tds_web")
FIG_DIR = BASE / "chuong4_figures"
DOCX_GLOB = "CH*.docx"


def setup_matplotlib_vn() -> None:
    for name in ("Segoe UI", "Arial", "Tahoma", "DejaVu Sans"):
        try:
            font_manager.findfont(name, fallback_to_default=False)
            plt.rcParams["font.family"] = name
            break
        except Exception:
            continue
    plt.rcParams["axes.unicode_minus"] = False


def save_fig(fig, name: str) -> Path:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    p = FIG_DIR / name
    fig.savefig(p, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return p


def figure_421_architecture() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    boxes = [
        (1, 3.5, "ESP32\nHTTP POST"),
        (4, 3.5, "Server\nPHP/XAMPP"),
        (7, 3.5, "MySQL\nsensor_data"),
        (7, 1.5, "Trình duyệt\nHTML/CSS/JS"),
    ]
    for x, y, t in boxes:
        ax.add_patch(
            mpatches.FancyBboxPatch(
                (x - 1, y - 0.65),
                2,
                1.3,
                boxstyle="round,pad=0.05",
                edgecolor="#2980b9",
                facecolor="#ecf5fc",
                linewidth=2,
            )
        )
        ax.text(x, y, t, ha="center", va="center", fontsize=11, wrap=True)

    ax.annotate("", xy=(3, 3.85), xytext=(3.95, 3.85), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.annotate("", xy=(6.05, 3.85), xytext=(5, 3.85), arrowprops=dict(arrowstyle="->", lw=1.8))
    ax.annotate("", xy=(7, 2.95), xytext=(7, 2.15), arrowprops=dict(arrowstyle="->", lw=1.8))

    ax.set_title(
        "Hình khối: luồng dữ liệu thời gian thực (ESP32 ↔ PHP ↔ CSDL)",
        fontsize=13,
        fontweight="bold",
        pad=12,
    )
    return save_fig(fig, "hinh421_khoi_he_thong.png")


def figure_422_database() -> Path:
    fig, ax = plt.subplots(figsize=(9, 3.8))
    ax.axis("off")
    ax.text(
        0.5,
        0.92,
        "Bảng sensor_data",
        fontsize=13,
        fontweight="bold",
        ha="center",
        transform=ax.transAxes,
    )
    rows = [
        "• id INT (PK, tự tăng)",
        "• time DATETIME",
        "• temp VARCHAR — nhiệt độ đã hiệu chỉnh",
        "• tds VARCHAR — TDS (ppm)",
        "• status VARCHAR — OK / LOW / HIGH",
        "• mode VARCHAR — non / truongthanh",
    ]
    y = 0.78
    for r in rows:
        ax.text(0.08, y, r, fontsize=12, transform=ax.transAxes)
        y -= 0.11

    ax.text(
        0.5,
        0.12,
        "Bảng app_mode — lưu chế độ cây cho ngưỡng cảnh báo TDS",
        fontsize=12,
        fontweight="bold",
        ha="center",
        transform=ax.transAxes,
    )
    return save_fig(fig, "hinh422_cosodu_lieu.png")


def figure_423_frontend() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    layers = [(5, 4.5, "Giao diện (HTML + CSS):\nHEADER, ô thông tin sinh viên,\nbảng dữ liệu, nút chức năng"),
              (5, 2.5, "Đồ thị Chart.js đọc JSON từ get-data.php"),
              (5, 0.8, "JavaScript: làm mới AJAX, đổi ngày, clear dữ liệu")]
    for x, y, t in layers:
        ax.add_patch(
            mpatches.FancyBboxPatch(
                (1, y - 0.72),
                8,
                1.44,
                boxstyle="round,pad=0.08",
                edgecolor="#27ae60",
                facecolor="#eafaf1",
                linewidth=2,
            )
        )
        ax.text(x, y, t, ha="center", va="center", fontsize=11)
    ax.annotate("", xy=(5, 4.1), xytext=(5, 3.42), arrowprops=dict(arrowstyle="->", lw=1.5))
    ax.annotate("", xy=(5, 2.07), xytext=(5, 1.54), arrowprops=dict(arrowstyle="->", lw=1.5))
    plt.suptitle("Tầng hiển thị web hydroponic TDS", fontsize=13, fontweight="bold", y=1.02)
    return save_fig(fig, "hinh423_giao_dien_web.png")


def figure_424_php() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.add_patch(
        mpatches.FancyBboxPatch(
            (0.5, 3.3),
            3,
            1.4,
            boxstyle="round,pad=0.05",
            edgecolor="#8e44ad",
            facecolor="#f4ecf7",
            linewidth=2,
        )
    )
    ax.text(2, 4, "ESP32 POST\n(insert.php)\nFields: temp, tds", ha="center", va="center", fontsize=10)
    ax.add_patch(
        mpatches.FancyBboxPatch(
            (6.5, 3.3),
            3,
            1.4,
            boxstyle="round,pad=0.05",
            edgecolor="#8e44ad",
            facecolor="#f4ecf7",
            linewidth=2,
        )
    )
    ax.text(8, 4, "config.php:\nkết nối mysqli +\nCREATE TABLE IF NOT EXISTS", ha="center", va="center", fontsize=10)
    ax.add_patch(
        mpatches.FancyBboxPatch(
            (2.9, 1.05),
            4.2,
            2.1,
            boxstyle="round,pad=0.08",
            edgecolor="#8e44ad",
            facecolor="#faf5fc",
            linewidth=2,
        )
    )
    ax.text(
        5,
        2.05,
        "insert.php nhận dữ liệu\n→ Kiểm tra API key (nếu cấu hình UPLOAD_API_KEY)"
        "\n→ Gán mode từ CSDL → Cảnh báo theo ngưỡng\n→ Chuỗi thời gian VN (+07)",
        ha="center",
        va="center",
        fontsize=10,
        linespacing=1.35,
    )
    ax.annotate("", xy=(4.6, 2.9), xytext=(2.4, 3.95), arrowprops=dict(arrowstyle="->", lw=1.4))
    ax.annotate("", xy=(5.4, 2.9), xytext=(7.7, 3.95), arrowprops=dict(arrowstyle="->", lw=1.4))

    arrows = [(2.2, 1.95, 3.95, 3.95), (7.85, 1.95, 6.1, 3.95)]
    for xa, ya, xb, yb in arrows:
        ax.annotate("", xy=(xb, yb), xytext=(xa, ya), arrowprops=dict(arrowstyle="->", lw=1.2))

    plt.suptitle(
        "Xử lý phía máy chủ: insert.php và cấu hình CSDL",
        fontsize=13,
        fontweight="bold",
        y=1.01,
    )
    return save_fig(fig, "hinh424_php_esp.png")


def figure_425_http() -> Path:
    fig, ax = plt.subplots(figsize=(9, 3))
    ax.axis("off")
    snippet = (
        "POST https://your-domain.com/insert.php\n"
        "Content-Type: application/x-www-form-urlencoded\n"
        "Header (tuỳ chọn): X-Api-Key: <UPLOAD_API_KEY>\n"
        "Body mẫu: temp=<float>&tds=<float>"
    )
    ax.text(
        0.5,
        0.52,
        snippet,
        ha="center",
        va="center",
        fontsize=10,
        transform=ax.transAxes,
        bbox=dict(boxstyle="round", facecolor="#fff9e6", edgecolor="#f39c12"),
    )
    ax.text(
        0.5,
        0.12,
        "ESP32 dùng thư viện HTTPClient; phản hồi OK/ERR/Unauthorized",
        ha="center",
        fontsize=12,
        transform=ax.transAxes,
    )
    plt.suptitle("Kết nối ESP32 qua HTTP", fontsize=13, fontweight="bold", y=1.06)
    return save_fig(fig, "hinh425_http_post.png")


def figure_426_realtime() -> Path:
    fig, ax = plt.subplots(figsize=(9, 2.9))
    ax.axis("off")
    tx = (
        "Trang index.php gọi get-data.php (JSON) định kỳ + get-mode.php\n"
        "→ Cập nhật bảng nhiệt độ / TDS / trạng thái màu (OK, LOW, HIGH)\n"
        "→ Đồng bộ với chế độ cây trên data.php / get-mode"
    )
    ax.text(0.5, 0.55, tx, ha="center", va="center", fontsize=12, transform=ax.transAxes)
    ax.add_patch(
        mpatches.FancyBboxPatch(
            (0.12, 0.12),
            0.76,
            0.28,
            boxstyle="round,pad=0.02",
            transform=ax.transAxes,
            facecolor="#ebf8ff",
            edgecolor="#3498db",
            linewidth=2,
        )
    )
    ax.text(
        0.5,
        0.26,
        "Ví dụ hiển thị: 26.4 °C  |  TDS 612 ppm  |  Trạng thái: OK",
        ha="center",
        fontsize=12,
        fontweight="bold",
        transform=ax.transAxes,
    )
    plt.suptitle("Dashboard thời gian thực", fontsize=13, fontweight="bold", y=1.05)
    return save_fig(fig, "hinh426_dashboard_realtime.png")


def figure_427_charts() -> Path:
    fig, ax = plt.subplots(figsize=(9, 3.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.add_patch(mpatches.Rectangle((1, 1.2), 3.5, 3, fill=True, facecolor="#fef9e7", edgecolor="#d4ac0d", lw=2))
    ax.add_patch(mpatches.Rectangle((5.5, 1.2), 3.5, 3, fill=True, facecolor="#eaf2f8", edgecolor="#2980b9", lw=2))
    ax.text(2.75, 3.5, "Biểu đồ\nnhiệt độ", ha="center", fontsize=11, fontweight="bold")
    ax.plot([1.8, 2.5, 3.3, 3.95], [2.8, 3.9, 2.5, 3.7], "-o", color="#c0392b", lw=2, markersize=5)
    ax.text(7.25, 3.5, "Biểu đồ\nTDS (ppm)", ha="center", fontsize=11, fontweight="bold")
    ax.plot([6.2, 6.9, 7.6, 8.35], [2.9, 2.6, 3.8, 3.45], "-o", color="#1f618d", lw=2, markersize=5)
    plt.suptitle("Chart.js nhận JSON và vẽ theo nhãn thời gian", fontsize=13, fontweight="bold", y=1.03)
    return save_fig(fig, "hinh427_charts_web.png")


def figure_428_history() -> Path:
    fig, ax = plt.subplots(figsize=(9, 2.9))
    ax.axis("off")
    ax.text(
        0.5,
        0.72,
        "Mỗi lần ESP32 POST → một bản ghi mới trong sensor_data",
        fontsize=13,
        fontweight="bold",
        ha="center",
        transform=ax.transAxes,
    )
    ax.text(
        0.5,
        0.38,
        "get-data.php: trả về 200 bản ghi mới nhất hoặc lọc theo ?date=YYYY-MM-DD\n"
        "export.php hỗ trợ xuất dữ liệu để báo cáo (nếu triển khai)",
        ha="center",
        fontsize=11,
        linespacing=1.38,
        transform=ax.transAxes,
    )
    plt.suptitle("Lịch sử và truy vấn", fontsize=13, fontweight="bold", y=1.05)
    return save_fig(fig, "hinh428_lich_su.png")


def figure_429_git() -> Path:
    fig, ax = plt.subplots(figsize=(9, 2.6))
    ax.axis("off")
    cmds = ("git clone / git pull → chỉnh sửa cục bộ\n"
            "git add . → git commit -m \"...\" → git push origin main\n"
            "GitHub là bản online; Railway triển khai từ repo")
    ax.text(
        0.5,
        0.55,
        cmds,
        ha="center",
        va="center",
        fontsize=12,
        bbox=dict(facecolor="#f6f8fa", edgecolor="#24292f", linewidth=1.6),
        transform=ax.transAxes,
    )
    ax.text(
        0.5,
        0.12,
        "(Ví dụ remote: hydroponic-tds-2 — chứa php, composer nếu có, readme)",
        ha="center",
        fontsize=10,
        style="italic",
        transform=ax.transAxes,
    )
    plt.suptitle("Quy trình đưa mã nguồn lên GitHub", fontsize=13, fontweight="bold", y=1.08)
    return save_fig(fig, "hinh429_github.png")


def figure_4210_deploy() -> Path:
    fig, ax = plt.subplots(figsize=(9, 3))
    ax.axis("off")
    txt = (
        "1. Tạo dịch vụ MySQL (hoặc dùng MySQL của host)\n"
        "2. Thiết lập biến môi trường: MYSQLHOST, MYSQLUSER,"
        " MYSQLPASSWORD,\n   MYSQLDATABASE, MYSQLPORT,"
        " UPLOAD_API_KEY (tuỳ chọn), APP_TIMEZONE\n"
        "3. Triển khai thư mục public chứa index.php và các PHP phụ\n"
        "4. Kiểm tra insert.php và index.php sau khi build"
    )
    ax.text(0.5, 0.52, txt, ha="center", va="center", fontsize=11, linespacing=1.42, transform=ax.transAxes)
    plt.suptitle("Triển khai lên Railway (PaaS)", fontsize=13, fontweight="bold", y=1.05)
    return save_fig(fig, "hinh4210_railway.png")


def figure_4211_final() -> Path:
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    ax.add_patch(
        mpatches.FancyBboxPatch(
            (0.4, 0.5),
            9.2,
            4.2,
            boxstyle="round,pad=0.05",
            edgecolor="#2c3e50",
            facecolor="#ffffff",
            linewidth=2.5,
        )
    )
    ax.add_patch(mpatches.Rectangle((0.7, 4.2), 8.6, 0.55, facecolor="#2980b9"))
    ax.text(5, 4.48, "Hydroponic System — Tiêu đề & thông tin đề tài", color="white", ha="center", fontsize=11, fontweight="bold")
    ax.add_patch(mpatches.Rectangle((0.8, 2.6), 4, 1.35, facecolor="#eef2f7", ec="#bdc3c7"))
    ax.add_patch(mpatches.Rectangle((5.2, 2.6), 4, 1.35, facecolor="#eef2f7", ec="#bdc3c7"))
    ax.text(2.8, 3.45, "Bảng dữ liệu", ha="center", fontsize=10, fontweight="bold")
    ax.text(7.2, 3.45, "Biểu đồ", ha="center", fontsize=10, fontweight="bold")
    ax.text(5, 1.6, "Nút: chọn ngày  |  Clear  |  Đổi chế độ cây", ha="center", fontsize=10)
    ax.text(5, 0.95, "Minh hoạ tổng hợp giao diện sau khi hoàn thiện", ha="center", fontsize=10, style="italic")
    plt.suptitle("Dashboard web sau khi hoàn thiện (wireframe)", fontsize=13, fontweight="bold", y=1.02)
    return save_fig(fig, "hinh4211_dashboard_hoan_thien.png")


def delete_paragraph(paragraph) -> None:
    """Remove paragraph from document body."""
    el = paragraph._element
    el.getparent().remove(el)


def remove_empty_between(doc: Document, start_text: str, end_text: str) -> None:
    """Strip empty paragraphs strictly between anchors (exclusive)."""

    paras = list(doc.paragraphs)
    i0 = next(i for i, p in enumerate(paras) if p.text.strip() == start_text)
    i1 = next(i for i, p in enumerate(paras) if p.text.strip() == end_text)
    for p in paras[i0 + 1 : i1]:
        if not p.text.strip():
            delete_paragraph(p)


def find_para(doc: Document, exact: str):
    for p in doc.paragraphs:
        if p.text.strip() == exact:
            return p
    raise ValueError(exact)


def add_figure_before(anchor, img: Path, caption: str, width_mm: float = 130) -> None:
    img_p = anchor.insert_paragraph_before()
    img_p.paragraph_format.space_after = Mm(1)
    img_p.add_run().add_picture(str(img), width=Mm(width_mm))
    anchor.insert_paragraph_before(caption)


def prepend_paragraphs_before_anchor(anchor, texts: list[str]) -> None:
    """Đưa các đoạn vào NGAY TRƯỚC anchor theo đúng thứ tự trong `texts` (trên→dưới)."""
    for line in texts:
        anchor.insert_paragraph_before(line)


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    setup_matplotlib_vn()

    imgs = {
        "4.2.1": figure_421_architecture(),
        "4.2.2": figure_422_database(),
        "4.2.3": figure_423_frontend(),
        "4.2.4": figure_424_php(),
        "4.2.5": figure_425_http(),
        "4.2.6": figure_426_realtime(),
        "4.2.7": figure_427_charts(),
        "4.2.8": figure_428_history(),
        "4.2.9": figure_429_git(),
        "4.2.10": figure_4210_deploy(),
        "4.2.11": figure_4211_final(),
    }

    doc_paths = sorted([p for p in BASE.glob(DOCX_GLOB) if not p.name.startswith("~$")])
    if not doc_paths:
        print("Không thấy file CHƯƠNG*.docx")
        return 1
    docx_path = doc_paths[0]

    headings = [
        ("4.2.1 Thiết kế cấu trúc tổng thể của hệ thống", "4.2.2 Thiết kế cơ sở dữ liệu của hệ thống"),
        ("4.2.2 Thiết kế cơ sở dữ liệu của hệ thống", "4.2.3 Xây dựng giao diện web bằng HTML, CSS và JavaScript"),
        (
            "4.2.3 Xây dựng giao diện web bằng HTML, CSS và JavaScript",
            "4.2.4 Xây dựng file PHP xử lý dữ liệu từ ESP32",
        ),
        ("4.2.4 Xây dựng file PHP xử lý dữ liệu từ ESP32", "4.2.5 Kết nối ESP32 với server thông qua giao thức HTTP"),
        (
            "4.2.5 Kết nối ESP32 với server thông qua giao thức HTTP",
            "4.2.6 Hiển thị dữ liệu thời gian thực trên web dashboard",
        ),
        (
            "4.2.6 Hiển thị dữ liệu thời gian thực trên web dashboard",
            "4.2.7 Xây dựng biểu đồ nhiệt độ và nồng độ dinh dưỡng trên web",
        ),
        (
            "4.2.7 Xây dựng biểu đồ nhiệt độ và nồng độ dinh dưỡng trên web",
            "4.2.8 Lưu trữ dữ liệu lịch sử trên server",
        ),
        ("4.2.8 Lưu trữ dữ liệu lịch sử trên server", "4.2.9 Đưa source code lên GitHub"),
        ("4.2.9 Đưa source code lên GitHub", "4.2.10 Triển khai web lên Railway để giám sát từ xa"),
        (
            "4.2.10 Triển khai web lên Railway để giám sát từ xa",
            "4.2.11 Kết quả giao diện web dashboard sau khi hoàn thiện",
        ),
        (
            "4.2.11 Kết quả giao diện web dashboard sau khi hoàn thiện",
            "4.3 Kết quả xây dựng ứng dụng Android",
        ),
    ]

    doc = Document(str(docx_path))
    remove_empty_between(
        doc,
        "4.2.11 Kết quả giao diện web dashboard sau khi hoàn thiện",
        "4.3 Kết quả xây dựng ứng dụng Android",
    )

    paras = doc.paragraphs
    anchor_42_main = None
    for p in paras:
        if p.text.strip() == "4.2 Kết quả xây dựng hệ thống web giám sát":
            anchor_42_main = p
            break
    if anchor_42_main is None:
        print("Không tìm thấy tiêu đề 4.2")
        return 1

    next_after_42 = find_para(doc, "4.2.1 Thiết kế cấu trúc tổng thể của hệ thống")
    intro_paras = [
        (
            "Hệ thống web giám sát thủy canh được tổ chức thành ba lớp: thiết bị đo ESP32,"
            " máy chủ XAMPP/PHP và cơ sở dữ liệu MySQL, và lớp hiển thị trên dashboard."
        ),
        (
            "Các phiên PHP (insert.php, get-data.php, data.php…) kết nối MySQL; trình duyệt nhận HTML và JSON qua fetch."
        ),
    ]

    prepend_paragraphs_before_anchor(next_after_42, intro_paras)

    SECTION_TEXT: dict[str, tuple[list[str], str]] = {
        "4.2.1": (
            [
                "Ở lớp thiết bị, firmware ESP32 gửi các cặp giá trị nhiệt độ và TDS đã được hiệu chỉnh theo POST HTTP."
                " Ở lớp máy chủ, thư viện HTTPClient của ESP32 chỉ vận chuyển gói dữ liệu; không cần MQTT hay socket riêng khi chỉ có ít nút và tần số POST vừa phải.",
                "Đường đi của dòng đo: ESP32 đo và kiểm tra kết nối Wi‑Fi → gửi POST tới địa chỉ insert.php của host"
                " (localhost trong quá trình phát triển, domain trong triển khai online)."
                " Máy chủ xác nhận, gán thời gian theo timezone cấu hình và chèn vào bảng sensor_data."
                " Dashboard đọc lại giá trị mới nhất qua các endpoint JSON.",
            ],
            "Hình 4.4. Sơ đồ khối luồng dữ liệu trong hệ thống web giám sát",
        ),
        "4.2.2": (
            [
                "Bảng app_mode chỉ chứa một dòng cố định (INSERT IGNORE) nhằm lưu chế độ cây: non hoặc truongthanh."
                " Các ngưỡng LOW/HIGH của trường status trong insert.php thay đổi tuỳ theo giá trị mode đọc từ CSDL.",
                "Khi không có máy chủ hoặc bảng mới được tạo lần đầu, lệnh CREATE TABLE IF NOT EXISTS trong config.php"
                " tự khởi tạo cấu trúc bảng sensor_data và app_mode trong MySQL InnoDB và charset utf8mb4.",
                "Để các trường temp và tds có thể lưu cả định dạng làm tròn từ thiết bị, kiểu VARCHAR(32)"
                " được giữ trong lược đồ thực tế của đề tài (có thể chuyển sang DECIMAL nếu cần bài toàn toán học thuần).",
            ],
            "Hình 4.5. Cấu trúc bảng dữ liệu chính trên máy chủ MySQL",
        ),
        "4.2.3": (
            [
                "Màu chữ và nền của header/dashboard được chỉnh bằng CSS (ví dụ nền #eef2f7, họ Segoe UI) để dễ đọc và thống nhất với bố cục in trong báo cáo.",
                "Biểu đồ được khởi tạo bằng Chart.js được tải từ CDN; JavaScript chia dữ liệu thành mảng thời gian và mảng thông số rồi gọi Chart.update để không phải tải lại toàn trang.",
                "Trang index.php gom một file HTML được nhúng trực tiếp (hoặc tách khối) với các thẻ có id để DOMContentLoaded và fetch gắn sự kiện.",
                "Đáp ứng thiết bị di động: các lưới thông tin sinh viên và bố cục header dùng @media nhỏ hơn 700 px để chuyển một cột.",
            ],
            "Hình 4.6. Tầng giao diện web (HTML, CSS, Chart.js và JavaScript)",
        ),
        "4.2.4": (
            [
                "Các endpoint gồm insert.php (ghi), get-latest.php và get-data.php (đọc), data.php và get-mode.php (quản trị/ghi nhận mode), export.php và delete.php (bao trì)."
                " Với mật chủ được đặt bằng biến môi trường UPLOAD_API_KEY thì chỉ client gửi đúng header X-Api-Key mới được phép INSERT.",
                "Đoạn mã trong insert.php dùng prepared statement và bind kiểu chuỗi cho các thông số, giảm nguy cơ SQL injection và lỗi ghép chuỗi.",
                "Riêng cờ APP_TIMEZONE mặc định là Asia/Ho_Chi_Minh để các mốc thời gian trong bản ghi thống nhất với giờ trong nước.",
            ],
            "Hình 4.7. Luồng tiếp nhận và lưu dữ liệu từ ESP32 qua các file PHP",
        ),
        "4.2.5": (
            [
                "Sau POST thành công, máy chủ có thể trả về chuỗi OK; trong lỗi CSDL có thể trả ERR với HTTP 500."
                " Mức 401 xuất hiện khi API key cấu hình nhưng ESP32 không gửi khóa khớp phía server.",
                "Trình tự được kiểm thử trong mạng LAN (XAMPP) trước: ESP32 và máy tính cùng LAN, URL trỏ tới máy chủ có IP LAN."
                " Sau đó chỉnh URL và biến môi trường Railway tương ứng khi đưa hệ lên máy chủ public.",
                "Nội dung body dùng form-urlencoded là lựa chọn đơn giản và tương thích PHP $_POST không cần JSON parser hai phía đồng bộ trong giai đoạn đầu.",
            ],
            "Hình 4.8. Minh họa định dạng yêu cầu HTTP POST từ ESP32 đến insert.php",
        ),
        "4.2.6": (
            [
                "Các ô số được cập nhật mà không tải lại trang nhờ hàm setInterval và fetch với Cache-Control và tham số t ngẫu nhiên tránh CDN cache.",
                "Trạng thái được mã hoá bằng màu theo LOW/OK/HIGH giúp người nuôi trông nhanh bằng mắt thường, giống trên cổng Arduino Serial Monitor.",
                "Chọn ngày thống nhất trong giao diện sẽ thêm query string dạng ?date=YYYY-MM-DD"
                " khiến get-data.php chỉ trả các bản ghi trong ngày, phục vụ xem nhật ký theo lịch làm việc.",
            ],
            "Hình 4.9. Cơ chế cập nhật dashboard thời gian thực",
        ),
        "4.2.7": (
            [
                "Đường nhiệt độ và đường TDS chia trục thời gian chung; khi không có điểm, biểu đồ không vẽ thêm và tránh báo lỗi JSON rỗng.",
                "Các thống kê và scale min/max của Chart được điều chỉnh theo nhãn và giá trị thực tế trả về, tránh các giá trị outlier ép biểu đồ không đọc được.",
                "Khi có nhiều hơn 200 điểm hiển thị, việc lọc theo ngày giúp biểu đồ gọn và phản ánh rõ xu hướng trong ca trực.",
            ],
            "Hình 4.10. Minh hoạ biểu đồ nhiệt độ và TDS trên web",
        ),
        "4.2.8": (
            [
                "Mỗi lần đo gửi lên máy chủ tạo một dòng mới (không gom 30 giây theo thiết kế hiện tại), nên có thể tính tốc độ lấy mẫu thực tế theo sai lệch thời gian giữa hai id liên tiếp.",
                "Tùy chọn clear (data.php với query clear=1) cho phép xoá và đặt lại lịch sử phục vụ các vòng thử mới không lẫn pha trước.",
                "Định danh AUTO_INCREMENT của id được dùng làm chỉ báo độ nhạy của cảnh báo theo chiều tăng dần của thử nghiệm.",
            ],
            "Hình 4.11. Lượt ghi lịch sử và lọc ngày trên máy chủ",
        ),
        "4.2.9": (
            [
                "Đưa nhánh main lên origin giúp đồng bộ các commit giữa máy chủ của sinh viên và bản được deploy (Railway thường build từ repo public).",
                "Tệp không cần theo Git (chuỗi kết nối cục bộ, cổng…) nhét vào gitignore và chỉ chia sẻ mẫu .env trong tài liệu hướng dẫn.",
                "Quyền collaborator trên GitHub giúp GVHD hoặc bạn học tái kiểm tra mã không cần copy USB.",
            ],
            "Hình 4.12. Quy trình làm việc với Git/GitHub trong đề tài",
        ),
        "4.2.10": (
            [
                "Sau deploy, có thể mở trang index và gửi thử một POST có temp/tds trong Postman để xác nhận chuỗi end‑to‑end trước khi ghép firmware thật.",
                "Biến môi trường không hard-code trong tái bản của config.php trong repo,"
                " nên chỉ chỉnh trên Railway dashboard hoặc dịch vụ host tương đương.",
                "Độ trễ mạng từ Railway về trong nước có thể lớn hơn LAN; trong báo cáo nên đo độ trễ trung bình rời rạc và ghi vào đánh giá hệ.",
            ],
            "Hình 4.13. Kiểm thử các bước triển khai trên nền tảng Railway",
        ),
        "4.2.11": (
            [
                "Khi được triển khai online và ESP32 được cấu hình địa chỉ đúng, người dùng thấy cùng bộ chỉ báo và biểu đồ trong trình duyệt không cần cài đặt bổ sung ngoại trình trên máy chủ học viện.",
                "Giai đoạn hoàn thiện có thể bổ sung logo trường, thông tin thành viên nhóm và bảng màu của theme đề tài (xanh navy #2980b9…) để thống nhất với bản báo cáo in.",
                "Hình minh hoạ dưới đây là wireframe tổng hợp; có thể thay bằng ảnh chụp màn hình thực tế từ máy local (http://localhost/tds_web/) hoặc domain production.",
            ],
            "Hình 4.14. Giao diện dashboard hoàn chỉnh (minh họa/ghi nhận kết quả)",
        ),
    }

    for cur, nxt in reversed(headings):
        key = cur.split()[0].strip()
        anchor = find_para(doc, nxt)
        texts, caption = SECTION_TEXT[key]
        prepend_paragraphs_before_anchor(anchor, texts)
        add_figure_before(anchor, imgs[key], caption, width_mm=138)

    out = docx_path
    try:
        doc.save(str(out))
    except PermissionError:
        alt = out.with_name(out.stem + " - bổ sung 4.2" + out.suffix)
        doc.save(str(alt))
        print("LƯU Ý: Không ghi được file gốc (có thể Word đang mở).")
        print("Đã lưu bản sao:", alt)
    else:
        print("Đã cập nhật:", docx_path)
    print("Hình trong:", FIG_DIR)
    return 0


if __name__ == "__main__":
    sys.exit(main())
